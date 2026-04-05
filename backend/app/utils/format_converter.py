import os
import tempfile
import zipfile
from io import BytesIO
from docx import Document
from PyPDF2 import PdfReader
from striprtf.striprtf import rtf_to_text
from pdf2docx import Converter
from fpdf import FPDF
import mammoth
from xhtml2pdf import pisa

# -----------------------------
# Utility helpers
# -----------------------------

def save_temp_file(file_bytes: bytes, suffix: str) -> str:
    """Saves bytes to a temporary file and returns the path."""
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
    tmp.write(file_bytes)
    tmp.close()
    return tmp.name

def read_file_to_bytes(path: str) -> BytesIO:
    """Reads a file into a BytesIO stream and deletes the file."""
    try:
        with open(path, "rb") as f:
            data = f.read()
        stream = BytesIO(data)
        stream.seek(0)
        return stream
    finally:
        if os.path.exists(path):
            os.remove(path)

def text_to_pdf_stream(text: str) -> BytesIO:
    """Helper to convert plain text to a PDF stream."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    # Multi_cell handles word wrapping
    pdf.multi_cell(0, 10, text)
    
    # Output to string/bytes
    # fpdf's output(dest='S') returns a latin1 encoded string in python 3
    pdf_bytes = pdf.output(dest='S').encode('latin1')
    return BytesIO(pdf_bytes)

def docx_to_pdf_stream(file_bytes: bytes) -> BytesIO:
    """
    Converts DOCX to PDF while preserving formatting (headings, bold, tables) 
    using mammoth (DOCX -> HTML) and xhtml2pdf (HTML -> PDF).
    """
    # 1. Convert DOCX to HTML
    result = mammoth.convert_to_html(BytesIO(file_bytes))
    html_content = result.value

    # 2. Wrap HTML in a template with basic CSS styling
    html_template = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <style>
            @page {{ size: A4; margin: 2cm; }}
            body {{ font-family: Helvetica, Arial, sans-serif; font-size: 11pt; line-height: 1.5; }}
            table {{ border-collapse: collapse; width: 100%; margin-bottom: 15px; }}
            th, td {{ border: 1px solid #000; padding: 6px; text-align: left; }}
            h1, h2, h3, h4 {{ color: #333333; }}
            p {{ margin-bottom: 10px; }}
        </style>
    </head>
    <body>
        {html_content}
    </body>
    </html>
    """

    # 3. Generate PDF from HTML
    pdf_stream = BytesIO()
    pisa_status = pisa.CreatePDF(BytesIO(html_template.encode('utf-8')), dest=pdf_stream)
    
    if pisa_status.err:
        raise ValueError("Failed to render PDF from generated HTML.")
        
    pdf_stream.seek(0)
    return pdf_stream

# -----------------------------
# MAIN UNIVERSAL CONVERTER
# -----------------------------

def convert_file(file_bytes: bytes, filename: str, target_format: str) -> BytesIO:
    source_ext = filename.split(".")[-1].lower()
    target_format = target_format.lower()

    if source_ext == target_format:
        return BytesIO(file_bytes)

    # Pre-validation: Ensure DOCX files are valid Zip archives to prevent crashes
    if source_ext == "docx":
        try:
            zipfile.ZipFile(BytesIO(file_bytes))
        except zipfile.BadZipFile:
            raise ValueError("Invalid DOCX file. It is corrupted or not a valid word document.")

    # 1. PDF -> DOCX (Preserves layout using pdf2docx)
    if source_ext == "pdf" and target_format == "docx":
        pdf_path = save_temp_file(file_bytes, ".pdf")
        docx_path = pdf_path.replace(".pdf", ".docx")
        try:
            cv = Converter(pdf_path)
            cv.convert(docx_path)
            cv.close()
            return read_file_to_bytes(docx_path)
        finally:
            if os.path.exists(pdf_path): os.remove(pdf_path)

    # 2. PDF -> TXT
    if source_ext == "pdf" and target_format == "txt":
        reader = PdfReader(BytesIO(file_bytes))
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
        return BytesIO(text.encode("utf-8"))

    # 3. DOCX -> TXT
    if source_ext == "docx" and target_format == "txt":
        doc = Document(BytesIO(file_bytes))
        text = "\n".join([para.text for para in doc.paragraphs])
        return BytesIO(text.encode("utf-8"))

    # 4. ANY -> PDF
    if target_format == "pdf":
        if source_ext == "docx":
            # Formatted conversion
            return docx_to_pdf_stream(file_bytes)
        elif source_ext == "rtf":
            text = rtf_to_text(file_bytes.decode("utf-8", errors="ignore"))
            return text_to_pdf_stream(text)
        elif source_ext == "txt":
            text = file_bytes.decode("utf-8", errors="ignore")
            return text_to_pdf_stream(text)
        else:
            raise ValueError(f"Cannot convert {source_ext} to PDF directly.")

    # 5. TXT/RTF -> DOCX
    if source_ext in ["txt", "rtf"] and target_format == "docx":
        doc = Document()
        if source_ext == "txt":
            text = file_bytes.decode("utf-8", errors="ignore")
        else:
            text = rtf_to_text(file_bytes.decode("utf-8", errors="ignore"))
        
        for line in text.split('\n'):
            doc.add_paragraph(line)

        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return output

    # 6. Fallback/Default error
    raise ValueError(f"Unsupported conversion: {source_ext} to {target_format}")