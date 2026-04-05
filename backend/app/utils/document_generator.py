from io import BytesIO
from docx import Document
from fpdf import FPDF

def create_docx(title: str, content: str) -> BytesIO:
    doc = Document()
    doc.add_heading(title, level=0)
    doc.add_paragraph(content)
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

def create_pdf(title: str, content: str) -> BytesIO:
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", 'B', 16)
    pdf.multi_cell(0, 10, title, align='C')
    pdf.set_font("Arial", '', 12)
    pdf.ln(5)
    pdf.multi_cell(0, 8, content)
    pdf_bytes = pdf.output(dest='S').encode('latin-1')
    file_stream = BytesIO(pdf_bytes)
    file_stream.seek(0)
    return file_stream